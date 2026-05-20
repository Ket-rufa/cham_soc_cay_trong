import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import json
import urllib.request

mermaid_code = """erDiagram
    users {
        bigint id PK
        varchar name
        varchar email
        timestamp email_verified_at
        varchar password
        varchar remember_token
        timestamp created_at
        timestamp updated_at
    }
    plants {
        bigint id PK
        bigint user_id FK
        varchar name
        varchar image_url
        varchar location
        text note
        timestamp created_at
        timestamp updated_at
    }
    plant_histories {
        bigint id PK
        bigint plant_id FK
        varchar action
        text note
        timestamp created_at
        timestamp updated_at
    }
    plant_care_schedules {
        bigint id PK
        bigint plant_id FK
        varchar task_type
        int frequency_days
        datetime last_done_at
        datetime next_due_at
        text note
        boolean is_active
        timestamp created_at
        timestamp updated_at
    }
    plant_libraries {
        bigint id PK
        varchar name
        varchar type
        text description
        text image_url
        varchar difficulty
        varchar light
        varchar water
        timestamp created_at
        timestamp updated_at
    }
    pest_disease_guides {
        bigint id PK
        varchar plant_name
        varchar disease_name
        varchar type
        text symptoms
        text causes
        text prevention
        text treatment
        text affected_plants
        json top_affected_flowers
        varchar image_url
        timestamp created_at
        timestamp updated_at
    }
    articles {
        bigint id PK
        varchar title
        text subtitle
        varchar category
        varchar read_time
        varchar icon_name
        varchar start_color
        varchar end_color
        json quick_tips
        json sections
        boolean is_published
        int sort_order
        timestamp created_at
        timestamp updated_at
    }

    users ||--o{ plants : "sở hữu"
    plants ||--o{ plant_histories : "có lịch sử"
    plants ||--o{ plant_care_schedules : "có lịch trình"
"""

state = {
    "code": mermaid_code,
    "mermaid": {"theme": "default"}
}
json_state = json.dumps(state)
b64_state = base64.urlsafe_b64encode(json_state.encode('utf-8')).decode('utf-8')
url = f"https://mermaid.ink/img/{b64_state}?type=png"

image_path = "mo_hinh_quan_he.png"

try:
    print(f"Downloading ER diagram...")
    req = urllib.request.Request(
        url, 
        data=None, 
        headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
    )
    with urllib.request.urlopen(req) as response, open(image_path, 'wb') as out_file:
        data = response.read()
        out_file.write(data)
    print("Downloaded successfully.")
except Exception as e:
    print(f"Error downloading image: {e}")

try:
    doc = docx.Document('Thiet_ke_co_so_du_lieu.docx')
    
    # Add new section
    doc.add_page_break()
    heading = doc.add_paragraph()
    run = heading.add_run('3.4.3. Mô hình quan hệ và vật lý')
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=Inches(6.0))
    
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_caption = p_caption.add_run('Hình 3.14: Mô hình quan hệ và vật lý của cơ sở dữ liệu')
    run_caption.italic = True
    run_caption.font.size = Pt(13)
    run_caption.font.name = 'Times New Roman'
    
    doc.save('Thiet_ke_co_so_du_lieu.docx')
    print("Updated Thiet_ke_co_so_du_lieu.docx successfully.")
except Exception as e:
    print(f"Error updating docx: {e}")
